import os
import json
from docx import Document

directory = "catalog/"

doc = Document()

for filename in os.listdir(directory):
    if filename.endswith(".json"):
        filepath = os.path.join(directory, filename)
        
        # Open the file and load the JSON data
        with open(filepath, 'r') as file:
            data = json.load(file)

        # Add a new paragraph for the filename
        doc.add_paragraph(f"Data from file: {filename}")
        
        # Add data to the doc
        doc.add_paragraph(f"Name: {data['name']}")
        doc.add_paragraph(f"Price: {data['price']}")
        doc.add_paragraph(f"Description: {data['description']}")

        # Add images if present
        if data.get('images'):
            images = ', '.join(data['images'])
            doc.add_paragraph(f"Images: {images}")

        # Add a line break after each file
        doc.add_paragraph("\n")

doc.save('page.docx')